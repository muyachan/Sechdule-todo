"""
Stage 5：產生完成報告（Word）

依《個人 PWA｜AI 開發交接摘要》第六節「期望交付成果」的十二個項目撰寫，
目標讀者是軟體新手，所以每個技術名詞都附白話說明，
並且把「你需要做什麼」放在最前面，而不是埋在文件末端。

刻意的設計決定：
  - 報告不判斷成敗，只如實呈現各階段結果。判斷由 Stage 4 負責。
  - 未執行的測試、已知限制一定要列出來，不能因為看起來不好看就省略。
"""

import json
import subprocess
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TAIPEI = timezone(timedelta(hours=8))


def _git(args: list) -> str:
    result = subprocess.run(["git"] + args, capture_output=True, text=True)
    return result.stdout.strip()


def collect_git_info(base_branch: str = "main") -> dict:
    return {
        "branch": _git(["rev-parse", "--abbrev-ref", "HEAD"]),
        "commit": _git(["rev-parse", "--short", "HEAD"]),
        "commit_message": _git(["log", "-1", "--pretty=%s"]),
        "changed_files": [
            f for f in _git(["diff", "--cached", "--name-only", base_branch]).split("\n") if f
        ],
    }


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_kv_table(doc, rows, col_widths=(Inches(2.0), Inches(4.2))):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    return table


def _add_bullets(doc, items, empty_text="（無）"):
    if not items:
        doc.add_paragraph(empty_text, style="List Bullet")
        return
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def build_report(
    report_data: dict,
    spec_title: str,
    output_path: str,
    pr_url: str = "",
    base_branch: str = "main",
):
    git_info = collect_git_info(base_branch)
    stage2 = report_data.get("stage2", {})
    stage3 = report_data.get("stage3", {})
    stage4 = report_data.get("stage4", {})

    now = datetime.now(TAIPEI).strftime("%Y-%m-%d %H:%M")
    meets = stage4.get("meets_spec", False)

    doc = Document()

    # ---------- 封面資訊 ----------
    title = doc.add_heading("自動協作管線｜完成報告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"任務：{spec_title}")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"產生時間：{now}（台北時間）").font.size = Pt(10)

    # ---------- 一、結論與你要做的事 ----------
    _add_heading(doc, "一、結論：你現在需要做什麼", level=1)

    p = doc.add_paragraph()
    run = p.add_run("整體判定：")
    run.bold = True
    verdict = run = p.add_run("通過審查，可以進行人工檢視" if meets else "未通過審查，需要你介入")
    verdict.bold = True
    verdict.font.color.rgb = RGBColor(0x1B, 0x7F, 0x3B) if meets else RGBColor(0xB3, 0x26, 0x1E)

    doc.add_paragraph(
        "這份報告由自動化管線產生。管線只負責「提出修改」與「初步審查」，"
        "所有實際生效的動作（合併、執行 SQL、部署）仍然只有你能做。"
    )

    _add_heading(doc, "合併前必須完成的步驟", level=2)
    todo_items = report_data.get("user_actions", [])
    if not todo_items:
        todo_items = [
            "打開下方的 Pull Request 連結，親自看過 diff 的每一行改動",
            "確認改動範圍沒有超出你原本要求的事",
            "若改動涉及前端檔案，記得 bump index.html 的 ?v= 與 sw.js 的 CACHE_NAME",
            "確認無誤後再按合併；有疑慮就直接關閉這個 PR，不會有任何副作用",
        ]
    _add_bullets(doc, todo_items)

    # ---------- 二、變更內容 ----------
    _add_heading(doc, "二、這次改了什麼", level=1)

    _add_kv_table(doc, [
        ("分支 branch", git_info["branch"] or "（尚未建立）"),
        ("提交 commit", git_info["commit"] or "（尚未提交）"),
        ("提交訊息", git_info["commit_message"] or "—"),
        ("Pull Request", pr_url or "（Stage 4 未通過，未建立 PR）"),
    ])

    doc.add_paragraph()
    _add_heading(doc, "修改的檔案清單", level=2)
    _add_bullets(doc, git_info["changed_files"], "（沒有任何檔案被修改）")

    # ---------- 三、各階段結果 ----------
    _add_heading(doc, "三、各階段執行結果", level=1)

    doc.add_paragraph(
        "管線分成三個自動階段。以下如實呈現每一段的結果，"
        "包含那些看起來不理想的部分。"
    )

    _add_heading(doc, "Stage 2：需求交叉檢查", level=2)
    doc.add_paragraph(
        "由 Claude 與 ChatGPT 各自審視需求規格，確認沒有機密外洩、程式崩潰、"
        "資料損失或違反架構限制的風險。預設放行，只有真正的問題才會擋下來。"
    )
    _add_kv_table(doc, [
        ("結果", "放行" if stage2.get("stop_reason") == "consensus_reached" else "被擋下"),
        ("進行輪數", stage2.get("final_round", 0)),
        ("關注關鍵字", "、".join(stage2.get("scope_warnings", [])) or "無"),
    ])

    doc.add_paragraph()
    _add_heading(doc, "Stage 3：Codex 實際修改程式", level=2)
    _add_kv_table(doc, [
        ("結果", "完成" if stage3.get("ok") else "失敗"),
        ("失敗原因", stage3.get("reason") or "—"),
    ])

    doc.add_paragraph()
    _add_heading(doc, "Stage 4：產出審查", level=2)
    doc.add_paragraph(
        "由 Claude 對照需求規格的驗收條件，逐條檢查 Codex 的改動是否達標。"
    )
    _add_kv_table(doc, [
        ("判定", "達標" if meets else "未達標"),
        ("說明", stage4.get("notes", "—")),
    ])

    if stage4.get("unmet_conditions"):
        doc.add_paragraph()
        _add_heading(doc, "未達成的驗收條件", level=3)
        _add_bullets(doc, stage4["unmet_conditions"])

    if stage4.get("violations"):
        doc.add_paragraph()
        _add_heading(doc, "違反的限制", level=3)
        _add_bullets(doc, stage4["violations"])

    # ---------- 四、測試 ----------
    _add_heading(doc, "四、測試狀況", level=1)

    _add_heading(doc, "已執行的測試", level=2)
    _add_bullets(doc, report_data.get("tests_run", []), "（本次未執行任何自動化測試）")

    _add_heading(doc, "未執行的測試（重要）", level=2)
    doc.add_paragraph(
        "以下項目自動化流程驗證不了，必須由你在實際裝置上確認。"
        "沒有做這些檢查，不代表功能正常，只代表還沒被檢查過。"
    )
    _add_bullets(doc, report_data.get("tests_not_run", []), "（無）")

    _add_heading(doc, "已知限制", level=2)
    _add_bullets(doc, report_data.get("known_limitations", []), "（無）")

    # ---------- 五、程式說明 ----------
    _add_heading(doc, "五、程式撰寫邏輯與函數說明", level=1)

    explanation = stage4.get("code_explanation") or {}
    overview = explanation.get("overview", "").strip()
    functions = explanation.get("functions") or []

    _add_heading(doc, "整體邏輯", level=2)
    doc.add_paragraph(overview or "（本次未產生程式說明）")

    if functions:
        doc.add_paragraph()
        _add_heading(doc, "各函數說明", level=2)
        for fn in functions:
            name = fn.get("name", "（未命名）")
            h = doc.add_paragraph()
            run = h.add_run(name)
            run.bold = True
            run.font.size = Pt(11)

            # how 欄位已移除（2026-08-02），改為只有 purpose 與 note，
            # 保留 how 的相容處理，讓舊報告資料也能正常呈現
            for label, key in (("負責什麼", "purpose"), ("怎麼做的", "how"), ("要注意什麼", "note")):
                value = (fn.get(key) or "").strip()
                if not value or value == "—":
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                lab = p.add_run(f"{label}：")
                lab.bold = True
                lab.font.size = Pt(10)
                val = p.add_run(value)
                val.font.size = Pt(10)
            doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph()
    note = p.add_run(
        "提醒：這份報告由 AI 產生，內容可能有誤。"
        "GitHub 上的實際 diff、Actions 執行紀錄與實機測試結果才是正式事實來源。"
    )
    note.italic = True
    note.font.size = Pt(9)

    doc.save(output_path)
    return output_path


def main():
    with open("pipeline_report.json", "r", encoding="utf-8") as f:
        report_data = json.load(f)

    import os
    pr_url = os.environ.get("PR_URL", "")
    spec_title = report_data.get("spec_path", "未命名任務")

    path = build_report(
        report_data=report_data,
        spec_title=spec_title,
        output_path="完成報告.docx",
        pr_url=pr_url,
    )
    print(f"[Stage 5] 報告已產生：{path}")


if __name__ == "__main__":
    main()
